from typing import List

from docx import Document
from docx.shared import Inches
from selenium.webdriver.common.by import By
from selenium.webdriver.remote.webdriver import WebDriver
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait

from puzzle import PopulatedPuzzle, Puzzle


class ReverseWordSearch(Puzzle):
    def __init__(self, grid_size_option: str):
        super().__init__()
        self.subdomain = "rws"
        self.grid_size_option = grid_size_option

    def get_options_internal(self, driver: WebDriver, timeout: int):
        input_gs = driver.find_element_by_id("sg")  # Grid Sizes
        gs_options = {
            size: f"{size}x{size}"
            for size in input_gs.get_attribute("data-slider-values").split(",")
        }

        return [("Grid Size", gs_options)]

    def select_options(self, driver: WebDriver, timeout: int):
        # For some reason, this doesn't work on the first puzzle
        input_gs = driver.find_element_by_id("sg")
        driver.execute_script(
            f"document.getElementById('sg').setAttribute('value', '{self.grid_size_option}')"
        )

    def get_populated_puzzle(self, driver: WebDriver, timeout: int):
        img_path = "tmp.png"

        board = (
            WebDriverWait(driver, timeout)
            .until(EC.presence_of_element_located((By.CLASS_NAME, "gridtable")))
            .find_element_by_tag_name("tbody")
        )
        board.screenshot(img_path)

        words = driver.find_element_by_class_name("wordlist").text.split("\n")
        return PopulatedReverseWordSearch(img_path, words)


class PopulatedReverseWordSearch(PopulatedPuzzle):
    def __init__(self, img: str, words: List[str], heading_style: str = "Heading 2"):
        super().__init__(heading_style)
        self.img = img
        self.words = words

    def add_to_doc_internal(self, doc: Document):
        doc.add_heading("Word List", level=2)
        doc.add_paragraph(", ".join(self.words))
        doc.add_picture(self.img, height=Inches(4))
        doc.add_page_break()
