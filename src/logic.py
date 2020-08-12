from docx import Document
from docx.shared import Inches
from selenium.webdriver.common.by import By
from selenium.webdriver.remote.webdriver import WebDriver
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select, WebDriverWait

from puzzle import PopulatedPuzzle, Puzzle


class LogicPuzzle(Puzzle):
    def __init__(self, grid_size_option: str, difficulty_option: str):
        super().__init__()
        self.subdomain = "logic"
        self.grid_size_option = grid_size_option
        self.difficulty_option = difficulty_option

    def get_options_internal(self, driver: WebDriver, timeout: int):
        select_gs = driver.find_element_by_id("sg")  # Grid Sizes
        gs_options = {
            option.get_attribute("value"): option.text
            for option in select_gs.find_elements_by_tag_name("option")
        }

        select_d = driver.find_element_by_id("sd")  # Difficulty
        d_options = {
            option.get_attribute("value"): option.text
            for option in select_d.find_elements_by_tag_name("option")
        }

        return [("Grid Size", gs_options), ("Difficulty", d_options)]

    def select_options(self, driver: WebDriver, timeout: int):
        select_gs = Select(driver.find_element_by_id("sg"))
        select_gs.select_by_value(self.grid_size_option)

        select_d = Select(driver.find_element_by_id("sd"))
        select_d.select_by_value(self.difficulty_option)

    def get_populated_puzzle(self, driver: WebDriver, timeout: int):
        img_path = "tmp.png"

        board = (
            WebDriverWait(driver, timeout)
            .until(EC.presence_of_element_located((By.ID, "puzzletable")))
            .find_element_by_tag_name("tbody")
        )
        board.screenshot(img_path)

        clues = "\n".join(
            clue.text
            for clue in driver.find_element_by_id("tabs-1").find_elements_by_class_name(
                "clue"
            )
        )

        driver.find_element_by_link_text("Story").click()
        story = (
            WebDriverWait(driver, timeout)
            .until(EC.visibility_of_element_located((By.ID, "tabs-2")))
            .find_element_by_tag_name("p")
            .text.strip()
        )

        return PopulatedLogicPuzzle(img_path, story, clues, self.grid_size_option)


class PopulatedLogicPuzzle(PopulatedPuzzle):
    def __init__(
        self,
        img: str,
        story: str,
        clues: str,
        grid_size_option: str,
        heading_style: str = "Heading 2",
    ):
        super().__init__(heading_style)
        self.img = img
        self.story = story
        self.clues = clues
        self.grid_size_option = grid_size_option

    def add_to_doc_internal(self, doc: Document):
        table = doc.add_table(2, 2)
        if self.is_small():
            table.cell(0, 0).merge(table.cell(0, 1))
            self.fill_cell(table.cell(0, 0), self.story, "Backstory and Goal")
            table.cell(1, 0).paragraphs[0].add_run().add_picture(
                self.img, height=Inches(3)
            )
            self.fill_cell(table.cell(1, 1), self.clues, "Clues")
        else:
            table.cell(1, 0).merge(table.cell(1, 1))
            table.cell(0, 0).paragraphs[0].add_run().add_picture(
                self.img, height=Inches(4)
            )
            self.fill_cell(table.cell(0, 1), self.story, "Backstory and Goal")
            self.fill_cell(table.cell(1, 0), self.clues, "Clues")

    def is_small(self):
        return True if int(self.grid_size_option) < 3 else False
