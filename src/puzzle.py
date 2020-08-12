import os
import sys

from docx import Document
from docx.document import Document as Doc
from docx.table import _Cell
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.desired_capabilities import DesiredCapabilities
from selenium.webdriver.remote.webdriver import WebDriver
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait


def get_driver():
    options = webdriver.ChromeOptions()
    options.headless = True
    options.add_experimental_option("excludeSwitches", ["enable-logging"])
    return webdriver.Remote(
        f"http://{os.environ.get('SELENIUM_HOST')}:{os.environ.get('SELENIUM_PORT')}/wd/hub",
        desired_capabilities=DesiredCapabilities.CHROME,
        options=options,
    )


class Puzzle:
    def get_options(self, driver: WebDriver, timeout: int):
        driver.get(f"https://{self.subdomain}.puzzlebaron.com/init.php")
        return self.get_options_internal(driver, timeout)

    def get_puzzle(self, driver: WebDriver, timeout: int):
        driver.get(f"https://{self.subdomain}.puzzlebaron.com/init.php")

        self.select_options(driver, timeout)

        # Create Puzzle
        driver.find_element_by_name("CreatePuzzle").click()
        WebDriverWait(driver, timeout).until(
            EC.presence_of_element_located((By.NAME, "submit"))
        ).click()

        # Set window size for screenshots
        driver.set_window_size(
            self.scroll_extent(driver, "Width"), self.scroll_extent(driver, "Height")
        )

        return self.get_populated_puzzle(driver, timeout)

    def scroll_extent(self, driver: WebDriver, direction: str) -> str:
        return driver.execute_script(
            f"return document.body.parentNode.scroll{direction}"
        )

    def get_options_internal(self, driver: WebDriver, timeout: int):
        pass

    def select_options(self, driver: WebDriver, timeout: int):
        pass

    def get_populated_puzzle(self, driver: WebDriver, timeout: int):
        pass


class PopulatedPuzzle:
    def __init__(self, heading_style: str = "Heading 2"):
        self.heading_style = "Heading 2"

    def fill_cell(self, cell: _Cell, text: str, heading: str = None):
        if heading is None:
            cell.paragraphs[0].text = text
        else:
            cell.paragraphs[0].style = self.heading_style
            cell.paragraphs[0].text = heading
            cell.add_paragraph(text)

    def add_to_doc(self, doc: Document):
        # print('Adding to document...')
        self.add_to_doc_internal(doc)
        doc.add_paragraph()

    def add_to_doc_internal(self, doc: Document):
        pass
