from docx import Document
from docx.shared import Inches
from selenium.webdriver.common.by import By
from selenium.webdriver.remote.webdriver import WebDriver
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait

from puzzle import PopulatedPuzzle, Puzzle


class Crossword(Puzzle):
    def __init__(self):
        super().__init__()
        self.subdomain = "crosswords"
        self.has_options = False

    def get_populated_puzzle(self, driver: WebDriver, timeout: int):
        img_path = "tmp.png"

        board = WebDriverWait(driver, timeout).until(
            EC.visibility_of_element_located((By.ID, "grid_container"))
        )
        board.screenshot(img_path)

        clues = driver.find_element_by_id("cluebox").text
        return PopulatedCrossword(img_path, clues)


class PopulatedCrossword(PopulatedPuzzle):
    def __init__(self, img: str, clues: str, heading_style: str = "Heading 2"):
        super().__init__(heading_style)
        self.img = img
        self.clues = clues

    def add_to_doc_internal(self, doc: Document):
        doc.add_picture(self.img, width=Inches(7))
        doc.add_heading("Clues", level=2)
        doc.add_paragraph(self.clues)
        doc.add_page_break()
